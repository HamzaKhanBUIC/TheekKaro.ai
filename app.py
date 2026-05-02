import streamlit as st
import google.generativeai as genai
import os
import time
import json
import re
import tempfile
import random

# --- Civic Features Imports ---
from docx import Document
from io import BytesIO
import urllib.parse
from streamlit_geolocation import streamlit_geolocation
from geopy.geocoders import Nominatim

MODEL_NAME = "gemini-2.5-flash-lite"

def get_api_keys():
    """Extracts valid 39-character Gemini API keys using Regex for maximum robustness.
    Priority: System Environment Variables (GCP Cloud Run) → Streamlit Secrets (local dev).
    The env var check runs first to avoid Streamlit's 'No secrets file found' warning on Cloud.
    """
    import re
    import os
    
    # 1. Try System Environment Variables FIRST (GCP Cloud Run — avoids secrets.toml warning)
    env_val = os.environ.get("GEMINI_API_KEYS", "")
    keys = re.findall(r'AIza[a-zA-Z0-9_-]{35}', env_val)
    if keys:
        return keys

    # 2. Fallback: Streamlit Secrets (local development only)
    # We read the value directly — if no secrets.toml exists (Cloud Run), this raises
    # a FileNotFoundError which we catch silently. No "in" check = no warning banner.
    try:
        raw_val = str(st.secrets.get("GEMINI_API_KEYS", ""))
        keys = re.findall(r'AIza[a-zA-Z0-9_-]{35}', raw_val)
        if keys:
            return keys
    except Exception:
        pass
    
    return []

# --- Page Configuration ---
st.set_page_config(
    page_title="Theek Karo AI — Reporting Citizens' Problems",
    page_icon="⚖️",  # Professional scales icon representing justice/civic duty
    layout="wide",
    initial_sidebar_state="collapsed"
)

# --- Premium Branding & Material Design CSS Injection ---
st.markdown("""
    <style>
        /* Import Google's Roboto Font & Font Awesome Icons */
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&family=Open+Sans:wght@400;600&display=swap');
        @import url('https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css');

        /* Root Palette & Typography Variables */
        :root {
            --primary-blue: #4285F4;
            --success-green: #34A853;
            --warning-yellow: #FBBC05;
            --bg-white: #FFFFFF;
            --text-dark: #1E293B; /* Darker navy for better contrast */
            --text-subtle: #475569;
            --border-light: #E2E8F0;
            --font-main: 'Roboto', 'Open Sans', 'Helvetica Neue', sans-serif;
            color-scheme: light; /* Tells browser to stay in light mode */
        }

        /* Global Structural Overrides */
        html, body, [data-testid="stAppViewContainer"] {
            background-color: var(--bg-white) !important;
            color: var(--text-dark) !important;
            font-family: var(--font-main) !important;
        }

        /* Remove Streamlit Visual Noise (Header, Hamburger, Footer) */
        header[data-testid="stHeader"] {
            visibility: hidden;
            height: 0% !important;
        }
        
        #MainMenu {
            visibility: hidden;
        }
        
        footer {
            visibility: hidden;
        }

        /* Fix Alert Box Contrast (st.info, st.success, etc.) */
        [data-testid="stNotification"] p {
            color: #1E293B !important;
            font-weight: 500 !important;
        }
        
        [data-testid="stNotification"] {
            border: 1px solid rgba(0,0,0,0.05) !important;
        }

        /* Adjusting Main Block Container for Breathing Room */
        .block-container {
            padding-top: 2rem !important;
            padding-bottom: 3rem !important;
            padding-left: 5rem !important;
            padding-right: 5rem !important;
            max-width: 1200px !important;
        }

        /* Premium Header Styling */
        .app-header {
            border-bottom: 1px solid var(--border-light);
            padding-bottom: 1.5rem;
            margin-bottom: 2.5rem;
        }
        
        .header-title {
            color: var(--text-dark) !important;
            font-size: 1.75rem;
            font-weight: 500;
            letter-spacing: -0.01em;
            margin-bottom: 0.2rem;
        }
        
        /* Force All Headings to be Dark */
        h1, h2, h3, h4, h5, h6 {
            color: var(--text-dark) !important;
            font-weight: 600 !important;
        }

        .header-descriptor {
            color: var(--text-subtle) !important;
            font-size: 0.85rem;
            font-weight: 400;
            letter-spacing: 0.02em;
            text-transform: uppercase;
        }

        /* Utility Classes for Material Design Elements */
        .stMarkdown p {
            color: var(--text-subtle) !important;
            line-height: 1.6;
        }

        /* Material Design Card Styling */
        .material-card {
            background-color: var(--bg-white);
            border: 1px solid var(--border-light);
            border-radius: 8px;
            padding: 1.5rem;
            transition: transform 0.2s ease, box-shadow 0.2s ease;
            margin-bottom: 1.5rem;
        }
        
        /* Force titles inside cards to be dark */
        .material-card h4 {
            color: var(--text-dark) !important;
            margin-bottom: 8px !important;
        }
        
        .material-card:hover {
            transform: translateY(-4px);
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            border-color: var(--primary-blue);
        }

        /* Hero Section Styling */
        .hero-container {
            position: relative;
            width: 100%;
            height: 300px;
            border-radius: 12px;
            overflow: hidden;
            margin-bottom: 2.5rem;
            background-color: #f8f9fa;
        }

        .hero-image {
            width: 100%;
            height: 100%;
            object-fit: cover;
            filter: brightness(0.6);
        }

        .hero-overlay {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            text-align: center;
            width: 80%;
            color: white !important;
        }

        .hero-title {
            font-size: 2.5rem !important;
            font-weight: 700 !important;
            margin-bottom: 0.5rem !important;
            color: white !important;
            text-shadow: 0 2px 4px rgba(0,0,0,0.3);
        }

        .hero-subtitle {
            font-size: 1.1rem !important;
            font-weight: 300 !important;
            color: rgba(255,255,255,0.9) !important;
        }

        /* Avatar Styling */
        .dev-avatar {
            border-radius: 12px;
            border: 1px solid var(--border-light);
            width: 100%;
            max-width: 160px;
            height: auto;
            display: block;
            margin: 0 auto 1rem auto;
        }

        /* --- Mobile Responsiveness --- */
        @media (max-width: 768px) {
            .block-container {
                padding-left: 1.2rem !important;
                padding-right: 1.2rem !important;
                padding-top: 1rem !important;
            }
            
            .hero-title {
                font-size: 1.5rem !important;
            }
            
            .hero-container {
                height: 180px;
            }

            .material-card {
                padding: 1.2rem;
            }
        }
    </style>
""", unsafe_allow_html=True)

# --- Phase 2: Hero Section ---
st.markdown("""
    <div class="hero-container">
        <img src="https://images.unsplash.com/photo-1517090504586-fde19ea6066f?auto=format&fit=crop&q=80&w=1200&h=300" class="hero-image">
        <div class="hero-overlay">
            <div class="hero-title">Transforming Complaints into Solutions</div>
            <div class="hero-subtitle">AI-powered civic reporting for a better tomorrow.</div>
        </div>
    </div>
""", unsafe_allow_html=True)

# --- Phase 2: Problem Type Cards ---
st.markdown("### Focus Areas")
p_col1, p_col2, p_col3 = st.columns(3)

with p_col1:
    st.markdown("""
        <div class="material-card">
            <h4 style="margin-top:0;">Infrastructure</h4>
            <p>Potholes, broken roads, and structural hazards in public spaces.</p>
        </div>
    """, unsafe_allow_html=True)

with p_col2:
    st.markdown("""
        <div class="material-card">
            <h4 style="margin-top:0;">Utilities</h4>
            <p>Faulty streetlights, open wiring, and electricity disruptions.</p>
        </div>
    """, unsafe_allow_html=True)

with p_col3:
    st.markdown("""
        <div class="material-card">
            <h4 style="margin-top:0;">Sanitation</h4>
            <p>Waste management issues, overflowing bins, and public hygiene.</p>
        </div>
    """, unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# --- Phase 2: About the Developer ---
with st.container():
    st.markdown('<div class="material-card">', unsafe_allow_html=True)
    st.markdown("### About the Developer")
    a_col1, a_col2 = st.columns([1, 2])
    
    with a_col1:
        # Responsive Avatar
        st.markdown('<img src="https://github.com/HamzaKhanBUIC/TheekKaro.ai/blob/main/dev_avatar.png?raw=true" class="dev-avatar">', unsafe_allow_html=True)
        # Fallback if raw image isn't available
        # st.image("dev_avatar.png", width=160) 
        
    with a_col2:
        st.markdown("""
            <div class="label-text">Lead Developer</div>
            <div class="detail-text">Hamza Imran</div>
            
            <div class="label-text">Institution</div>
            <div class="detail-text">Air University Islamabad</div>
            
            <div class="label-text">Academic Path</div>
            <div class="detail-text">BS Cyber Security (CYS)</div>
            
            <div class="label-text">Current Standing</div>
            <div class="detail-text">2nd Semester</div>

            <div class="social-links">
                <a href="https://www.linkedin.com/in/hamza-imran-17569b383" target="_blank" class="social-icon" title="LinkedIn">
                    <i class="fab fa-linkedin-in"></i>
                </a>
                <a href="https://github.com/HamzaKhanBUIC" target="_blank" class="social-icon" title="GitHub">
                    <i class="fab fa-github"></i>
                </a>
                <a href="mailto:hamza135252@gmail.com" class="social-icon" title="Email">
                    <i class="fas fa-envelope"></i>
                </a>
            </div>
        """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# --- Legal Disclaimer Banner ---
st.markdown("""
    <div style="
        background: linear-gradient(135deg, #FFF8E1, #FFF3E0);
        border-left: 4px solid #FF8F00;
        border-radius: 6px;
        padding: 1rem 1.5rem;
        margin: 1rem 0 1.5rem 0;
        font-family: var(--font-main);
    ">
        <p style="font-weight: 600; color: #E65100; margin: 0 0 0.4rem 0; font-size: 0.85rem; text-transform: uppercase; letter-spacing: 0.05em;">⚖️ Privacy & Legal Disclaimer</p>
        <p style="color: #4E342E; font-size: 0.82rem; line-height: 1.6; margin: 0;">
            <strong>TheekKaro.ai</strong> is a civic reporting tool built for public good. By using this platform, you acknowledge:
            <br>• <strong>No personal data is stored.</strong> Media files are processed transiently via Google Gemini API and deleted immediately after analysis.
            <br>• <strong>We are not responsible</strong> for the accuracy of authority contact information, which is AI-generated and may require verification.
            <br>• Complaints generated are for your personal use. We do not transmit, forward, or share your data with any government body on your behalf.
            <br>• Use of this tool constitutes acceptance of <a href="https://policies.google.com/privacy" target="_blank" style="color: #E65100;">Google's Privacy Policy</a> (Gemini API) and our fair-use terms.
            <br>• This platform is a student hackathon project and carries <strong>no official government affiliation.</strong>
        </p>
    </div>
""", unsafe_allow_html=True)

st.markdown("<br><hr><br>", unsafe_allow_html=True)
# ---------------------------------
# --- Multimodal Input ---
st.markdown("### Capture or Upload Hazard")
st.info("**Pro-Tip:** If the network is slow, uploading a **Picture** is much faster and more reliable than a video.")

# Two columns for Video vs Camera
tab1, tab2 = st.tabs(["Initiate Secure Media Upload", "Initialize Camera Node"])

with tab1:
    uploaded_file = st.file_uploader(
        "Select hazard evidence for analysis", 
        type=["mp4", "mov", "jpg", "jpeg", "png"]
    )

with tab2:
    camera_file = st.camera_input("Capture live hazard data")

# Prioritize camera if used, otherwise file uploader
final_file = camera_file if camera_file is not None else uploaded_file

def repair_node(key):
    """Self-healing function: Wipes all files from a node if it hits quota limits."""
    try:
        import google.generativeai as genai
        genai.configure(api_key=key)
        files = genai.list_files()
        for f in files:
            f.delete()
        return True
    except:
        return False

# --- UI Logic ---
if final_file is not None:
    # 1. Determine file type
    is_video = final_file.type in ["video/mp4", "video/quicktime"]
    suffix = ".mp4" if is_video else ".jpg"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        tmp_file.write(final_file.read())
        tmp_media_path = tmp_file.name

    # --- Real GPS Detection & Reverse Geocoding ---
    st.markdown("### Detect Location")
    st.info("Please allow location access so we can route the complaint accurately.")
    location = streamlit_geolocation(key='theekkaro_loc_unique')

    if location and location.get('latitude') is not None and location.get('longitude') is not None:
        lat = location['latitude']
        lon = location['longitude']
        
        with st.spinner("Translating GPS coordinates into street address..."):
            try:
                geolocator = Nominatim(user_agent="TheekKaro_Hackathon_App")
                location_data = geolocator.reverse(f"{lat}, {lon}")
                human_address = location_data.address if location_data else "Unknown Address"
                user_location = f"{human_address} (GPS: {lat}, {lon})"
                st.success(f"Location locked: {human_address}")
            except Exception:
                user_location = f"Latitude: {lat}, Longitude: {lon}"
                st.success("GPS coordinates locked!")
    else:
        user_location = "Karachi, Pakistan"
        
    # --- Refined AI Prompt ---
    prompt = (
        f"You are an expert civic assistant AI helping Pakistani citizens report public hazards to the correct government authority. "
        f"A citizen has submitted visual evidence of a civic hazard from this location: {user_location}.\n\n"
        "TASK: Carefully analyze the image/video and produce a structured report. "
        "You MUST output ONLY a valid raw JSON object — absolutely NO markdown, NO code fences, NO extra commentary before or after the JSON.\n\n"
        "The JSON object must contain EXACTLY these 6 keys:\n"
        "1. 'hazard_detected': A precise 3-8 word description of the specific hazard visible (e.g., 'Collapsed sewage manhole cover', 'Exposed high-voltage wiring').\n"
        "2. 'severity': Exactly one of: 'Low', 'Medium', 'High', or 'Critical'. Base this on immediate public danger risk.\n"
        "3. 'relevant_authority': The EXACT name of the responsible Pakistani local government department (e.g., 'Karachi Water & Sewerage Board', 'LESCO Lahore Electric Supply Company', 'KMC Roads Department'). Be specific to the city/region inferred from the location.\n"
        "4. 'authority_email': The real, publicly listed official email address for that department. If genuinely unknown, use the string 'Not Found'. Do NOT invent or guess emails.\n"
        "5. 'english_complaint': A formal, professional complaint letter in English. It must include: (a) today's date, (b) the citizen's location coordinates/address, (c) a clear description of the hazard, (d) an urgent call to action with a requested deadline of 7 days, and (e) a closing statement signed 'A Concerned Citizen via TheekKaro.ai'.\n"
        "6. 'urdu_complaint': The same formal complaint letter fully translated and written in proper Urdu script. It must mirror all details from the English version.\n\n"
        "CRITICAL RULES: Output ONLY the JSON. No preamble. No explanation. The JSON must be parseable by Python's json.loads()."
    )

    # --- CONSOLIDATED PROCESSING ENGINE ---
    try:

            


        # --- CONSOLIDATED PROCESSING ENGINE (Upload + Analysis) ---
        st.info("Analyzing footage and generating official documents...")
        
        api_keys = get_api_keys()
        if not api_keys:
            st.error("No API Keys configured. Please set GEMINI_API_KEYS.")
            st.stop()
            
        random.shuffle(api_keys)
        response_text = None
        media_file = None
        
        # We try each key for the ENTIRE sequence (Upload -> Wait -> Analyze)
        for attempt, current_key in enumerate(api_keys):
            try:
                genai.configure(api_key=current_key)
                
                # 1. Upload
                media_file = genai.upload_file(path=tmp_media_path)
                
                # 2. Wait for processing (if video)
                if is_video:
                    timeout_seconds = 60
                    start_time = time.time()
                    while media_file.state.name == "PROCESSING":
                        if time.time() - start_time > timeout_seconds:
                            raise TimeoutError("Processing timeout.")
                        time.sleep(2)
                        media_file = genai.get_file(media_file.name)
                    
                    if media_file.state.name == "FAILED":
                        raise ValueError("File processing failed.")

                # 3. Analyze
                model = genai.GenerativeModel(MODEL_NAME)
                response = model.generate_content([media_file, prompt])
                response_text = response.text
                
                if response_text:
                    st.success(f"{'Video' if is_video else 'Image'} processed successfully!")
                    break # Sequence COMPLETE
                    
            except Exception as e:
                error_msg = str(e).lower()
                # SELF-HEALING: If node is busy/full, clean it up before switching
                if "429" in error_msg or "quota" in error_msg or "limit" in error_msg:
                    st.warning(f"Node {attempt + 1} exhausted. Initiating auto-repair...")
                    repair_node(current_key) # Wipe old data from this key
                    
                    if attempt < len(api_keys) - 1:
                        st.warning("Routing to secondary node...")
                        continue 
                    else:
                        st.error("All processing nodes are currently at maximum capacity. Please try again in 2 minutes.")
                        st.stop()
                else:
                    # Specific error reporting for internal node failures
                    st.error(f"Node failure: {str(e)}")
                    st.stop()

        if not response_text:
            st.error("Could not complete analysis. Please try a smaller file.")
            st.stop()

        raw_text = response_text
            
        # --- Error-Proof Parsing ---
        clean_json_str = raw_text.strip()
        if clean_json_str.startswith("```"):
            clean_json_str = re.sub(r"^```(?:json)?\n", "", clean_json_str)
            clean_json_str = re.sub(r"\n```$", "", clean_json_str)
            
        try:
            data = json.loads(clean_json_str)
            hazard = data.get("hazard_detected", "Unknown Hazard")
            severity = data.get("severity", "Unknown")
            authority = data.get("relevant_authority", "Relevant Authorities")
            authority_email = data.get("authority_email", "Not Found")
            english_text = data.get("english_complaint", "Content missing.")
            urdu_text = data.get("urdu_complaint", "Content missing.")
            
        except json.JSONDecodeError:
            st.warning("Received malformed formatting from the AI. Engaging data salvage protocols...")
            hazard = "Civic Hazard (Salvaged)"
            severity = "Unknown"
            authority = "Local Authorities"
            authority_email = "Not Found"
            
            eng_match = re.search(r'"english_complaint"\s*:\s*"(.*?)"(?=\s*,|\s*\})', raw_text, re.DOTALL)
            english_text = eng_match.group(1).replace('\\n', '\n').replace('\\"', '"') if eng_match else "Could not salvage English text."
            
            urdu_match = re.search(r'"urdu_complaint"\s*:\s*"(.*?)"(?=\s*,|\s*\})', raw_text, re.DOTALL)
            urdu_text = urdu_match.group(1).replace('\\n', '\n').replace('\\"', '"') if urdu_match else "Could not salvage Urdu text."

        # --- UI Display: Intelligence Dashboard ---
        st.markdown("---")
        st.markdown(f'<p style="font-family: var(--font-mono); color: #10B981; font-weight: bold; font-size: 1.1rem; margin-bottom: 1.5rem;">[ STATUS: HAZARD DETECTED ]</p>', unsafe_allow_html=True)

        # Logic for Severity Badging
        sev_color = "#10B981" # Default Green
        if severity in ["High", "Critical"]:
            sev_color = "#EF4444" # Red
        elif severity == "Medium":
            sev_color = "#F59E0B" # Amber

        # Metrics Grid: Custom Cards
        m_col1, m_col2, m_col3 = st.columns(3)
        
        with m_col1:
            st.markdown(f"""
                <div style="background-color: #1E293B; border-radius: 8px; padding: 20px; border: 1px solid #334155;">
                    <p style="color: #94A3B8; margin: 0; font-size: 0.75rem; text-transform: uppercase; font-weight: 600;">Detected Hazard</p>
                    <p style="color: #FFFFFF; margin: 0; font-size: 1.25rem; font-weight: 600; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;">{hazard}</p>
                </div>
            """, unsafe_allow_html=True)

        with m_col2:
            st.markdown(f"""
                <div style="background-color: #1E293B; border-radius: 8px; padding: 20px; border: 1px solid #334155;">
                    <p style="color: #94A3B8; margin: 0; font-size: 0.75rem; text-transform: uppercase; font-weight: 600;">Severity Level</p>
                    <p style="color: {sev_color}; margin: 0; font-size: 1.25rem; font-weight: 600;">{severity}</p>
                </div>
            """, unsafe_allow_html=True)

        with m_col3:
            st.markdown(f"""
                <div style="background-color: #1E293B; border-radius: 8px; padding: 20px; border: 1px solid #334155;">
                    <p style="color: #94A3B8; margin: 0; font-size: 0.75rem; text-transform: uppercase; font-weight: 600;">Routing Authority</p>
                    <p style="color: #FFFFFF; margin: 0; font-size: 1.25rem; font-weight: 600; white-space: nowrap; overflow: hidden; text-overflow: ellipsis;">{authority}</p>
                </div>
            """, unsafe_allow_html=True)
            
        st.markdown("<div style='margin-bottom: 2rem;'></div>", unsafe_allow_html=True)
            
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("English Complaint")
            st.text_area("To be sent to relevant officials:", value=english_text, height=400)
            
        with col2:
            st.subheader("Urdu Complaint (اردو)")
            st.text_area("متعلقہ حکام کے نام:", value=urdu_text, height=400)

        # --- UI Display: Action Buttons ---
        st.markdown("---")
        st.subheader("Take Action")
        action_col1, action_col2 = st.columns(2)

        with action_col1:
            if authority_email.lower() == "not found":
                st.warning(f"No public email found for {authority}. Search manually for: '{authority} official email Pakistan'")
            else:
                subject_enc = urllib.parse.quote(f"Urgent Civic Hazard Report: {hazard}")
                # Truncate body to 1800 chars to stay within URL limits on all browsers/platforms
                body_preview = english_text[:1800] + ("\n\n[Full complaint in attached .docx]" if len(english_text) > 1800 else "")
                body_enc = urllib.parse.quote(body_preview)
                mailto_link = f"mailto:{authority_email}?subject={subject_enc}&body={body_enc}"
                
                st.markdown(f"""
                    <div style="margin-bottom: 0.5rem;">
                        <a href="{mailto_link}" 
                           style="display:inline-block; background-color:#4285F4; color:white !important;
                                  padding: 0.6rem 1.2rem; border-radius:6px; text-decoration:none;
                                  font-weight:600; font-size:0.9rem; transition: background 0.2s;"
                           onmouseover="this.style.backgroundColor='#1A73E8'"
                           onmouseout="this.style.backgroundColor='#4285F4'">
                            ✉️ Send Email to Authority
                        </a>
                    </div>
                    <p style="font-size:0.75rem; color:#5F6368; margin-top:0.3rem;">
                        Sending to: <strong>{authority_email}</strong><br>
                        <em>If Gmail doesn't open automatically, copy the address above and paste into your email client.</em>
                    </p>
                """, unsafe_allow_html=True)
                
                # Fallback: Let user copy email address manually
                st.code(authority_email, language=None)

        with action_col2:
            doc = Document()
            doc.add_heading(f'Civic Hazard Report: {hazard}', 0)
            doc.add_heading('English Application', level=1)
            doc.add_paragraph(english_text)
            doc.add_heading('Urdu Application (اردو)', level=1)
            doc.add_paragraph(urdu_text)
            
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.download_button(
                label="EXPORT: .DOCX Payload",
                data=doc_buffer,
                file_name=f"TheekKaro_Complaint_{hazard.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"An error occurred during generation: {str(e)}")
        
    finally:
        if 'media_file' in locals() and media_file:
            try:
                genai.delete_file(media_file.name)
            except:
                pass
        try:
            os.remove(tmp_media_path) 
        except:
            pass